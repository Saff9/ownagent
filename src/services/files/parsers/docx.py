"""
DOCX parser for GenZ Smart
Extracts text from Word documents
"""
from typing import Dict, Any, Optional
from pathlib import Path


class DocxParser:
    """Parser for Word documents"""
    
    def __init__(self):
        self.name = "DOCX Parser"
        self.supported_extensions = [".docx"]
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX file and extract text
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        result = {
            "text": "",
            "metadata": {},
            "paragraphs": 0,
            "error": None
        }
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            result["metadata"] = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraphs": len(doc.paragraphs)
            }
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            result["paragraphs"] = len(text_parts)
            result["text"] = "\n\n".join(text_parts)
            result["word_count"] = len(result["text"].split())
            
        except ImportError:
            result["error"] = "python-docx not installed. Install with: pip install python-docx"
        except Exception as e:
            result["error"] = f"Failed to parse DOCX: {str(e)}"
        
        return result
    
    def get_preview(self, file_path: str, max_chars: int = 1000) -> str:
        """
        Get a preview of the document content
        
        Args:
            file_path: Path to DOCX file
            max_chars: Maximum characters to return
            
        Returns:
            Preview text
        """
        result = self.parse(file_path)
        
        if result["error"]:
            return f"Error: {result['error']}"
        
        text = result["text"]
        if len(text) > max_chars:
            text = text[:max_chars] + "..."
        
        return text
    
    def extract_headers(self, file_path: str) -> list:
        """
        Extract headers from document
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of header texts
        """
        headers = []
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            for para in doc.paragraphs:
                # Check if paragraph is a heading
                if para.style and para.style.name and para.style.name.startswith('Heading'):
                    headers.append(para.text)
                    
        except ImportError:
            pass
        
        return headers
